from pathlib import Path
import shutil

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
NEW_TITLE = (
    "Knowledge-Enhanced Large Language Models for Bilingual Railway Vocational "
    "Education under Resource Constraints"
)
TARGETS = (
    ROOT / "paper/ijwis/submission/anonymous_manuscript.docx",
    ROOT / "paper/ijwis/submission/title_page_template.docx",
)


def iter_paragraphs(document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def replace_title(document) -> int:
    title_paragraph = next(
        (
            paragraph
            for paragraph in iter_paragraphs(document)
            if paragraph.style.name in {"Title", "Heading 1"} and paragraph.text.strip()
        ),
        None,
    )
    if title_paragraph is None:
        return 0
    if len(title_paragraph.runs) == 1:
        title_paragraph.runs[0].text = NEW_TITLE
    else:
        title_paragraph.text = NEW_TITLE
    replacements = 1
    document.core_properties.title = NEW_TITLE
    return replacements


for target in TARGETS:
    backup = target.with_suffix(target.suffix + ".title-backup")
    shutil.copy2(target, backup)
    document = Document(target)
    count = replace_title(document)
    if count != 1:
        raise RuntimeError(f"Expected one title replacement in {target}, found {count}")
    document.save(target)
    print(f"updated {target}")
